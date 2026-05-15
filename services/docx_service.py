print("[docx_service.py] Loading DOCX service...")

import io
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── PAGE SETUP ───────────────────────────────────────────────────────────────
# Nigerian university standard: Times New Roman 12pt, double-spaced
# Margins: 1.5 inch left, 1 inch right/top/bottom

LEFT_MARGIN   = Inches(1.5)
RIGHT_MARGIN  = Inches(1.0)
TOP_MARGIN    = Inches(1.0)
BOTTOM_MARGIN = Inches(1.0)

FONT_NAME = "Times New Roman"
FONT_SIZE = 12


# ─── DOCUMENT SETUP ───────────────────────────────────────────────────────────

def _create_document() -> Document:
    """Create a base document with correct page setup."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin   = LEFT_MARGIN
        section.right_margin  = RIGHT_MARGIN
        section.top_margin    = TOP_MARGIN
        section.bottom_margin = BOTTOM_MARGIN

    # Default style — Times New Roman 12pt
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE)

    # Set default paragraph spacing — double spaced
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_before      = Pt(0)
    style.paragraph_format.space_after       = Pt(0)

    return doc


def _set_font(run, bold=False, italic=False, size=12):
    run.font.name  = FONT_NAME
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic


def _add_paragraph(doc, text="", bold=False, italic=False,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   size=12, space_before=0, space_after=0,
                   first_line_indent=None, keep_with_next=False) -> object:
    """Add a paragraph with full formatting control."""
    p   = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment          = align
    fmt.line_spacing_rule  = WD_LINE_SPACING.DOUBLE
    fmt.space_before       = Pt(space_before)
    fmt.space_after        = Pt(space_after)
    fmt.keep_with_next     = keep_with_next
    if first_line_indent is not None:
        fmt.first_line_indent = Inches(first_line_indent)

    if text:
        run = p.add_run(text)
        _set_font(run, bold=bold, italic=italic, size=size)

    return p


def _add_page_break(doc):
    doc.add_page_break()


def _add_chapter_heading(doc, chapter_number: int, chapter_name: str):
    """Bold centered ALL CAPS chapter heading matching Nigerian university standard."""
    # e.g. "CHAPTER ONE"
    p = _add_paragraph(
        doc,
        text=f"CHAPTER {_number_to_word(chapter_number)}",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0,
        space_after=0,
    )
    # e.g. "1.0 INTRODUCTION"
    _add_paragraph(
        doc,
        text=f"{chapter_number}.0 {chapter_name.upper()}",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0,
        space_after=12,
    )


def _add_section_heading(doc, text: str):
    """Bold left-aligned numbered section heading."""
    _add_paragraph(
        doc,
        text=text,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=12,
        space_after=6,
    )


def _add_body_paragraph(doc, text: str, first_in_section=False):
    """Body text — double spaced, justified, first line indent."""
    indent = None if first_in_section else 0.5
    _add_paragraph(
        doc,
        text=text,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=indent,
    )


def _number_to_word(n: int) -> str:
    words = {1: "ONE", 2: "TWO", 3: "THREE", 4: "FOUR", 5: "FIVE"}
    return words.get(n, str(n))


# ─── PRELIMINARY PAGES ────────────────────────────────────────────────────────

def _build_title_pages(doc, project: dict, user: dict):
    print("[docx] Building title pages...")
    university   = project.get("university", "Nigerian University").upper()
    department   = project.get("department", "").upper()
    faculty      = project.get("faculty", "").upper()
    topic        = project.get("topic", "").upper()
    level        = (project.get("academic_level") or "bsc").upper()
    student_name = (user.get("first_name", "Student") if user else "Student").upper()
    month_year   = datetime.now().strftime("%B, %Y").upper()

    degree_map = {
        "BSC": "BACHELOR OF SCIENCE",
        "HND": "HIGHER NATIONAL DIPLOMA",
        "PGD": "POSTGRADUATE DIPLOMA",
        "MSC": "MASTER OF SCIENCE",
        "MBA": "MASTER OF BUSINESS ADMINISTRATION",
        "MPA": "MASTER OF PUBLIC ADMINISTRATION",
        "PHD": "DOCTOR OF PHILOSOPHY",
    }
    degree = degree_map.get(level, "BACHELOR OF SCIENCE")

    # ── Cover page ────────────────────────────────────────────────────────────
    _add_paragraph(doc, space_before=36)  # top spacing
    _add_paragraph(doc, text=topic, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    _add_paragraph(doc, space_before=12)
    _add_paragraph(doc, text="BY", bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, text=student_name, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)
    if department:
        _add_paragraph(doc, text=f"DEPARTMENT OF {department}",
                       align=WD_ALIGN_PARAGRAPH.CENTER)
    if faculty:
        _add_paragraph(doc, text=faculty,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, text=university,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    _add_paragraph(doc, text=month_year,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_page_break(doc)

    # ── Submission page ───────────────────────────────────────────────────────
    _add_paragraph(doc, space_before=36)
    _add_paragraph(doc, text=topic, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    _add_paragraph(doc, space_before=12)
    _add_paragraph(doc, text="BY", bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, text=student_name, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    submission = (
        f"A PROJECT REPORT SUBMITTED TO THE DEPARTMENT OF "
        f"{department if department else 'THE DEPARTMENT'}, "
        f"{university}, IN PARTIAL FULFILLMENT FOR THE AWARD OF "
        f"{degree} DEGREE IN "
        f"{department if department else 'THE RELEVANT FIELD'}."
    )
    _add_paragraph(doc, text=submission,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)
    _add_paragraph(doc, text=month_year,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_page_break(doc)


def _build_declaration(doc, project: dict, user: dict):
    print("[docx] Building declaration...")
    topic        = project.get("topic", "")
    student_name = (user.get("first_name", "Student") if user else "Student").upper()

    _add_paragraph(doc, text="DECLARATION", bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    _add_body_paragraph(
        doc,
        f'I hereby declare that the study titled: "{topic}" is a collection of my '
        f"original research work, and it has not been presented for any other "
        f"qualification anywhere. Information from other sources (published and "
        f"unpublished) has been duly acknowledged.",
        first_in_section=True,
    )
    _add_paragraph(doc, space_before=36)
    _add_paragraph(doc, text="...............................................",
                   align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(doc, text=student_name, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_page_break(doc)


def _build_certification(doc, project: dict, user: dict):
    print("[docx] Building certification...")
    topic        = project.get("topic", "")
    student_name = (user.get("first_name", "Student") if user else "Student").upper()
    university   = project.get("university", "the University")
    department   = project.get("department", "the Department")
    level        = (project.get("academic_level") or "bsc").upper()
    degree_map   = {
        "BSC": "Bachelor of Science", "MSC": "Master of Science",
        "MBA": "MBA", "PHD": "Doctor of Philosophy",
    }
    degree = degree_map.get(level, "Bachelor of Science")

    _add_paragraph(doc, text="CERTIFICATION", bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    _add_body_paragraph(
        doc,
        f'This project report, entitled "{topic}" by {student_name}, '
        f"meets the regulations governing the award of the degree of {degree} "
        f"in {department} of {university}, and is approved for its contribution "
        f"to knowledge and literary presentation.",
        first_in_section=True,
    )

    _add_paragraph(doc, space_before=24)
    for label in ["Project Supervisor", "Head of Department", "External Examiner"]:
        _add_paragraph(doc, text="...............................................",
                       align=WD_ALIGN_PARAGRAPH.LEFT)
        _add_paragraph(doc, text=label, align=WD_ALIGN_PARAGRAPH.LEFT)
        _add_paragraph(doc, space_before=12)
    _add_page_break(doc)


def _build_dedication(doc):
    print("[docx] Building dedication...")
    _add_paragraph(doc, text="DEDICATION", bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    _add_body_paragraph(
        doc,
        "This project is dedicated to my family, whose unwavering support and "
        "encouragement made this work possible, and to all Nigerian students "
        "striving for academic excellence.",
        first_in_section=True,
    )
    _add_page_break(doc)


def _build_acknowledgement(doc, user: dict):
    print("[docx] Building acknowledgement...")
    _add_paragraph(doc, text="ACKNOWLEDGEMENT", bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    _add_body_paragraph(
        doc,
        "I extend my heartfelt gratitude to God Almighty, whose guidance has seen "
        "me through the successful completion of this project. I also thank my "
        "supervisor, my head of department, and all the lecturers who contributed "
        "to my academic journey. Special appreciation goes to my family and friends "
        "for their continuous support and encouragement throughout this programme.",
        first_in_section=True,
    )
    _add_page_break(doc)


def _build_abstract(doc, project: dict):
    print("[docx] Building abstract...")
    _add_paragraph(doc, text="ABSTRACT", bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    ch1 = project.get("chapter_1_content", "")
    if ch1:
        # Extract the first clean paragraphs up to ~300 words
        paragraphs = [p.strip() for p in ch1.split("\n") if p.strip()
                      and not p.strip().startswith("#")
                      and not p.strip().startswith("<!--")]
        words_count = 0
        abstract_parts = []
        for para in paragraphs[2:]:  # Skip chapter heading lines
            words = para.split()
            if words_count + len(words) > 300:
                break
            abstract_parts.append(para)
            words_count += len(words)
        abstract_text = " ".join(abstract_parts) if abstract_parts else ch1[:800]
        _add_body_paragraph(doc, _strip_markdown(abstract_text), first_in_section=True)
    else:
        _add_body_paragraph(
            doc,
            "[Abstract will appear here after Chapter 1 is generated and reviewed by the student.]",
            first_in_section=True,
        )

    topic_words = project.get("topic", "").split()[:6]
    keywords    = ", ".join(topic_words)
    _add_paragraph(doc, space_before=12)
    kw = doc.add_paragraph()
    kw.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_run = kw.add_run("Keywords: ")
    _set_font(kw_run, bold=True)
    kw_val = kw.add_run(keywords)
    _set_font(kw_val)
    _add_page_break(doc)


def _build_toc(doc, project: dict):
    print("[docx] Building table of contents...")
    from utils.constants import CHAPTER_NAMES

    _add_paragraph(doc, text="TABLE OF CONTENTS", bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    prelim = [
        ("Title page", "i"),
        ("Declaration", "iii"),
        ("Certification", "iv"),
        ("Dedication", "v"),
        ("Acknowledgement", "vi"),
        ("Abstract", "vii"),
        ("Table of Contents", "viii"),
        ("List of Figures", "xi"),
    ]
    for label, page in prelim:
        _add_toc_line(doc, label, page, bold=False)

    chapters_done = project.get("chapters_completed", 0)
    for ch_num in range(1, 6):
        ch_name   = CHAPTER_NAMES.get(ch_num, f"Chapter {ch_num}")
        page_ref  = str(ch_num) if ch_num <= chapters_done else "-"
        _add_toc_line(doc, f"CHAPTER {ch_num}: {ch_name.upper()}", page_ref, bold=True)

    _add_toc_line(doc, "REFERENCES", "-", bold=True)
    _add_page_break(doc)


def _add_toc_line(doc, label: str, page: str, bold=False):
    """Add a TOC line with dot leader and page number."""
    p   = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    fmt.space_before      = Pt(0)
    fmt.space_after       = Pt(0)

    label_run = p.add_run(label)
    _set_font(label_run, bold=bold)

    dots_run = p.add_run(f"{'.' * max(1, 60 - len(label))} {page}")
    _set_font(dots_run, bold=bold)


# ─── CHAPTER CONTENT ──────────────────────────────────────────────────────────

def _build_chapter(doc, chapter_number: int, content: str):
    print(f"[docx] Building chapter {chapter_number}...")
    from utils.constants import CHAPTER_NAMES
    chapter_name = CHAPTER_NAMES.get(chapter_number, f"Chapter {chapter_number}")

    _add_chapter_heading(doc, chapter_number, chapter_name)

    lines          = content.split("\n")
    first_in_section = True

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Section heading detection
        if _is_section_heading(line):
            clean = re.sub(r'^#{1,6}\s*', '', line).strip()
            _add_section_heading(doc, _strip_markdown(clean))
            first_in_section = True
            continue

        # Skip markdown table separator rows
        if re.match(r'^\|[-| :]+\|$', line):
            continue

        # Table rows — render as plain text for now
        if line.startswith("|") and line.endswith("|"):
            _add_paragraph(doc, text=_strip_markdown(line),
                           align=WD_ALIGN_PARAGRAPH.LEFT)
            first_in_section = False
            continue

        # Skip extracted data blocks
        if "<!--EXTRACTED_DATA" in line or "EXTRACTED_DATA-->" in line:
            continue

        # Body paragraph
        clean = _strip_markdown(line)
        if clean:
            _add_body_paragraph(doc, clean, first_in_section=first_in_section)
            first_in_section = False

    _add_page_break(doc)


def _is_section_heading(line: str) -> bool:
    if re.match(r'^#{1,6}\s+', line.strip()):
        return True
    if re.match(r'^\d+\.\d+(\.\d+)?\s+\S', line.strip()):
        return True
    return False


def _strip_markdown(text: str) -> str:
    """Remove markdown formatting for plain Word output."""
    text = re.sub(r'^#{1,6}\s+', '', text.strip())
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',    r'\1', text)
    text = re.sub(r'_(.+?)_',      r'\1', text)
    text = re.sub(r'`(.+?)`',      r'\1', text)
    return text.strip()


# ─── REFERENCE LIST ───────────────────────────────────────────────────────────

def _build_references(doc, project: dict):
    print("[docx] Building references...")
    from utils.constants import REFERENCE_HEADINGS

    citation_style = project.get("citation_style", "apa7") or "apa7"
    heading        = REFERENCE_HEADINGS.get(citation_style, "References").upper()

    _add_paragraph(doc, text=heading, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    refs_raw = project.get("verified_references", [])
    if isinstance(refs_raw, str):
        try:
            refs_raw = json.loads(refs_raw)
        except Exception:
            refs_raw = []

    if not refs_raw:
        _add_body_paragraph(
            doc,
            "[References will appear here after Chapter 2 is generated.]",
            first_in_section=True,
        )
        _add_page_break(doc)
        return

    # Sort alphabetically for author-date styles
    if citation_style not in ("ieee", "vancouver"):
        refs_raw = sorted(
            refs_raw,
            key=lambda r: (
                r.get("authors", [""])[0].split()[-1].lower()
                if r.get("authors") else "zzz"
            ),
        )

    for i, ref in enumerate(refs_raw, 1):
        formatted = _format_reference(ref, citation_style, i)
        if formatted:
            # Hanging indent for references
            p   = doc.add_paragraph()
            fmt = p.paragraph_format
            fmt.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.line_spacing_rule  = WD_LINE_SPACING.DOUBLE
            fmt.left_indent        = Inches(0.5)
            fmt.first_line_indent  = Inches(-0.5)
            fmt.space_before       = Pt(0)
            fmt.space_after        = Pt(6)
            run = p.add_run(formatted)
            _set_font(run)

    _add_page_break(doc)


def _format_reference(ref: dict, style: str, number: int) -> str:
    authors  = ref.get("authors", [])
    year     = ref.get("year", "n.d.")
    title    = ref.get("title", "Unknown Title")
    journal  = ref.get("journal", "")
    doi      = ref.get("doi", "")
    doi_url  = f"https://doi.org/{doi}" if doi else ""

    def apa_authors(authors):
        if not authors: return "Unknown Author"
        formatted = []
        for name in authors[:6]:
            parts = name.strip().split()
            if len(parts) >= 2:
                surname  = parts[-1]
                initials = ". ".join(p[0].upper() for p in parts[:-1]) + "."
                formatted.append(f"{surname}, {initials}")
            else:
                formatted.append(name)
        if len(authors) > 6:
            formatted.append("et al.")
        if len(formatted) > 1:
            return ", ".join(formatted[:-1]) + ", & " + formatted[-1]
        return formatted[0]

    def ieee_authors(authors):
        if not authors: return "Unknown"
        formatted = []
        for name in authors[:6]:
            parts = name.strip().split()
            if len(parts) >= 2:
                initials = ". ".join(p[0].upper() for p in parts[:-1]) + "."
                surname  = parts[-1]
                formatted.append(f"{initials} {surname}")
            else:
                formatted.append(name)
        if len(authors) > 6:
            formatted.append("et al.")
        return ", ".join(formatted)

    j = f"{journal}. " if journal else ""
    d = f" {doi_url}" if doi_url else ""

    if style == "apa7":
        return f"{apa_authors(authors)} ({year}). {title}. {j}{d}".strip()
    elif style == "harvard":
        return f"{apa_authors(authors)} ({year}) '{title}', {j}{d}".strip()
    elif style == "ieee":
        return f"[{number}] {ieee_authors(authors)}, '{title},' {j}{year}.{d}".strip()
    elif style == "vancouver":
        return f"{number}. {ieee_authors(authors)}. {title}. {j}{year}.{d}".strip()
    else:
        return f"{apa_authors(authors)} ({year}). {title}. {j}{d}".strip()


# ─── DISCLAIMER ───────────────────────────────────────────────────────────────

def _build_disclaimer(doc):
    _add_paragraph(doc, space_before=24)
    _add_paragraph(doc, text="AI ASSISTANCE DISCLOSURE", bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "This project was prepared with the assistance of FYP Mentor, an AI research assistant. "
        "The student is responsible for reviewing, verifying, editing, and taking full ownership "
        "of all content before submission. All citations should be verified against published "
        "academic sources. Review with your supervisor before submission."
    )
    run.font.name   = FONT_NAME
    run.font.size   = Pt(10)
    run.font.italic = True


# ─── MAIN GENERATION ──────────────────────────────────────────────────────────

def generate_project_docx(project: dict, user: dict = None) -> io.BytesIO:
    """
    Generate the complete project as a .docx file in memory.

    Format matches Nigerian university standard:
    - Times New Roman 12pt
    - Double spaced
    - 1.5 inch left margin, 1 inch other margins
    - Proper preliminary pages
    - Chapter headings: bold centered ALL CAPS
    - Section headings: bold left aligned numbered
    - Body text: justified with first line indent
    - References with hanging indent sorted alphabetically

    If the student provided a chapter format or per-chapter outline during onboarding,
    those are already baked into the chapter content by Claude.
    This function just formats whatever content is stored on the project record.

    Never writes to disk — returns BytesIO buffer.
    """
    print(f"[docx] generate_project_docx: project_id={project.get('id')}")

    # Merge user fields into project
    if user:
        project = dict(project)
        project["university"]     = project.get("university")     or user.get("university", "")
        project["department"]     = project.get("department")     or user.get("department", "")
        project["academic_level"] = project.get("academic_level") or user.get("academic_level", "bsc")
        project["faculty"]        = project.get("faculty")        or user.get("faculty", "")

    doc = _create_document()

    # Preliminary pages
    _build_title_pages(doc, project, user or {})
    _build_declaration(doc, project, user or {})
    _build_certification(doc, project, user or {})
    _build_dedication(doc)
    _build_acknowledgement(doc, user or {})
    _build_abstract(doc, project)
    _build_toc(doc, project)

    # Chapters
    chapters_done = project.get("chapters_completed", 0)
    for ch_num in range(1, chapters_done + 1):
        content = project.get(f"chapter_{ch_num}_content", "")
        if content:
            _build_chapter(doc, ch_num, content)
            print(f"[docx] Chapter {ch_num} added.")

    # References
    _build_references(doc, project)

    # Disclaimer
    _build_disclaimer(doc)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    size_kb = len(buffer.getvalue()) // 1024
    print(f"[docx] DOCX generated. Size: {size_kb} KB")
    return buffer


print("[docx_service.py] DOCX service loaded.")